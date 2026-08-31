#!/usr/bin/env python3
"""Render the progress-report Markdown subset as a Google-Docs-friendly DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INLINE_RE = re.compile(r"(\*\*.*?\*\*|`.*?`)")
IMAGE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")


def add_inline(paragraph, text: str) -> None:
    """Add the bold and inline-code subset used by the report."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(45, 45, 45)
        else:
            paragraph.add_run(part)


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    add_inline(paragraph, text.strip())
    if bold:
        for run in paragraph.runs:
            run.bold = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, lines: list[str]) -> None:
    parsed = [
        [item.strip() for item in line.strip().strip("|").split("|")] for line in lines
    ]
    rows = [parsed[0], *parsed[2:]]
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.start_type = WD_SECTION.NEW_PAGE
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (
        ("Title", 22, RGBColor(20, 52, 82)),
        ("Heading 1", 16, RGBColor(20, 72, 110)),
        ("Heading 2", 13, RGBColor(30, 86, 126)),
        ("Heading 3", 11, RGBColor(45, 90, 120)),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.keep_with_next = True


def render(markdown: Path, output: Path) -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = "M0 Second Cross-Country Federated Training Report"
    document.core_properties.subject = (
        "Federated training and GlobalOpinionQA evaluation"
    )
    document.core_properties.author = "Monash M0 collaboration team"

    lines = markdown.read_text(encoding="utf-8").splitlines()
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(item.strip() for item in paragraph_buffer))
        paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            flush_paragraph()
            caption, relative = image_match.groups()
            image_path = (markdown.parent / relative).resolve()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(image_path), width=Inches(6.85))
            caption_paragraph = document.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_paragraph.add_run(caption)
            caption_run.italic = True
            caption_run.font.size = Pt(9)
            index += 1
            continue
        if stripped.startswith("# "):
            flush_paragraph()
            document.add_heading(stripped[2:], level=0)
            index += 1
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            document.add_heading(stripped[3:], level=1)
            index += 1
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            document.add_heading(stripped[4:], level=2)
            index += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue
        numbered_match = NUMBERED_RE.match(stripped)
        if numbered_match:
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, numbered_match.group(2))
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    render(args.markdown.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
