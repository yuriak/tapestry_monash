#!/usr/bin/env python3
"""Build figures and DOCX for the M0 local-FL progress report."""

from __future__ import annotations

import argparse
import csv
import json
import re
from collections.abc import Iterable
from pathlib import Path
from typing import Any

SITES = ("au", "india")
SITE_LABELS = {"au": "Australia", "india": "India"}
COLORS = {"au": "#0072B2", "india": "#D55E00"}


def csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def make_plots(run_root: Path, assets: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import pandas as pd

    assets.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update(
        {
            "figure.dpi": 150,
            "savefig.dpi": 180,
            "font.size": 10,
            "axes.grid": True,
            "grid.alpha": 0.25,
            "axes.spines.top": False,
            "axes.spines.right": False,
        }
    )

    fig, ax = plt.subplots(figsize=(8.2, 4.5))
    for site in SITES:
        frame = pd.read_csv(run_root / site / "logs/epoch_loss_tracking_muon.csv")
        frame["global_step"] = range(1, len(frame) + 1)
        frame["progress"] = 100 * frame["global_step"] / len(frame)
        frame["rolling"] = frame["loss"].rolling(50, min_periods=10).mean()
        ax.plot(
            frame["progress"],
            frame["rolling"],
            color=COLORS[site],
            linewidth=1.8,
            label=SITE_LABELS[site],
        )
    ax.axvline(50, color="#666666", linestyle="--", linewidth=1)
    ax.text(0.51, 0.93, "Pass 2", color="#555555", transform=ax.transAxes)
    ax.set(xlabel="Normalized local training progress (%)", ylabel="Rolling-50 loss")
    ax.set_title("Local training loss across ten federated rounds")
    ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(assets / "training_loss_normalized.png", bbox_inches="tight")
    plt.close(fig)

    fig, axes = plt.subplots(2, 1, figsize=(9.0, 5.8), sharex=True)
    first_time = None
    frames = {}
    for site in SITES:
        frame = pd.read_csv(
            run_root / site / "gpu.csv",
            header=None,
            names=("timestamp", "gpu", "name", "util", "used", "total", "power"),
            skipinitialspace=True,
        )
        frame["timestamp"] = pd.to_datetime(frame["timestamp"])
        frames[site] = frame
        start = frame["timestamp"].min()
        first_time = start if first_time is None else min(first_time, start)
    for ax, site in zip(axes, SITES):
        frame = frames[site].sort_values("timestamp")
        frame["elapsed"] = (frame["timestamp"] - first_time).dt.total_seconds() / 60
        frame["rolling"] = frame["util"].rolling(60, min_periods=1).mean()
        ax.plot(frame["elapsed"], frame["rolling"], color=COLORS[site], linewidth=1)
        boundaries = []
        with (run_root / site / "logs/runtime_comm.log").open() as handle:
            next(handle)
            for line in handle:
                fields = line.rstrip().split(",")
                if len(fields) >= 4 and fields[3] == "node_start":
                    timestamp = pd.to_datetime(fields[0])
                    boundaries.append((timestamp - first_time).total_seconds() / 60)
        for boundary in boundaries:
            ax.axvline(boundary, color="#999999", linewidth=0.5, alpha=0.35)
        ax.set_ylim(0, 103)
        ax.set_ylabel(f"{SITE_LABELS[site]} GPU util. (%)")
    axes[0].set_title("Compute and wait cycles under fixed 13-minute federation windows")
    axes[-1].set_xlabel("Elapsed time (minutes)")
    fig.tight_layout()
    fig.savefig(assets / "gpu_utilization_timeline.png", bbox_inches="tight")
    plt.close(fig)

    identities = {
        site: json.loads((run_root / site / "identity.json").read_text())["node_id"]
        for site in SITES
    }
    fig, axes = plt.subplots(1, 2, figsize=(9.0, 3.9))
    for site in SITES:
        delta_rows = [
            json.loads(line)
            for line in (run_root / site / "logs/local_delta_audit.jsonl")
            .read_text()
            .splitlines()
        ]
        axes[0].plot(
            [row["round"] for row in delta_rows],
            [row["l2_norm"] for row in delta_rows],
            marker="o",
            color=COLORS[site],
            label=SITE_LABELS[site],
        )
        trust = csv_rows(run_root / site / "logs/trust_scores_new.csv")
        timestamps = sorted({row["timestamp"] for row in trust})
        peer_id = next(value for key, value in identities.items() if key != site)
        peer_weights = [
            float(
                next(
                    (row["weight"]
                    for row in trust
                    if row["timestamp"] == timestamp
                    and row["peer_node"] == peer_id),
                    0.0,
                )
            )
            for timestamp in timestamps
        ]
        axes[1].plot(
            range(1, len(peer_weights) + 1),
            peer_weights,
            marker="o",
            color=COLORS[site],
            label=f"{SITE_LABELS[site]} observer",
        )
    axes[0].set(title="Local LoRA-delta magnitude", xlabel="Round", ylabel="L2 norm")
    axes[1].set(
        title="Weight assigned to the other site",
        xlabel="Round",
        ylabel="Observer-local trust weight",
        ylim=(0, 1),
    )
    for ax in axes:
        ax.set_xticks(range(1, 11))
        ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(assets / "federation_dynamics.png", bbox_inches="tight")
    plt.close(fig)

    evaluation = run_root / "evaluation/full-grid"
    cultural = pd.read_csv(evaluation / "culturalbench/summary.csv")
    cultural = cultural[cultural["region_group"] == "overall"]
    goqa = pd.read_csv(evaluation / "global-opinion-qa-five-prompt/summary.csv")
    goqa = goqa[(goqa["source"] == "all") & (goqa["region_group"] == "all")]
    metrics = (
        ("easy_accuracy", "Easy accuracy", True),
        ("hard_binary_accuracy", "Hard binary accuracy", True),
        ("hard_question_exact_match", "Hard exact match", True),
        ("goqa", "GOQA country-macro JSD", False),
    )
    fig, axes = plt.subplots(2, 2, figsize=(9.2, 6.8))
    for ax, (column, title, percentage) in zip(axes.flat, metrics):
        base = (
            float(cultural.loc[cultural["checkpoint"] == "base", column].iloc[0])
            if column != "goqa"
            else float(goqa.loc[goqa["run"] == "base", "country_macro_mean_js_distance"].iloc[0])
        )
        ax.axhline(
            100 * base if percentage else base,
            color="#222222",
            linestyle="--",
            linewidth=1.2,
            label="Base",
        )
        for site in SITES:
            names = [f"fl_{site}_p{value:03d}" for value in (20, 40, 60, 80, 100)]
            if column == "goqa":
                values = [
                    float(
                        goqa.loc[
                            goqa["run"] == name, "country_macro_mean_js_distance"
                        ].iloc[0]
                    )
                    for name in names
                ]
                raw = float(
                    goqa.loc[
                        goqa["run"] == f"fl_{site}_p100_raw",
                        "country_macro_mean_js_distance",
                    ].iloc[0]
                )
            else:
                values = [
                    float(cultural.loc[cultural["checkpoint"] == name, column].iloc[0])
                    for name in names
                ]
                raw = float(
                    cultural.loc[
                        cultural["checkpoint"] == f"fl_{site}_p100_raw", column
                    ].iloc[0]
                )
            if percentage:
                values = [100 * value for value in values]
                raw *= 100
            ax.plot(
                (20, 40, 60, 80, 100),
                values,
                color=COLORS[site],
                marker="o",
                linewidth=1.8,
                label=SITE_LABELS[site],
            )
            ax.scatter(
                [100],
                [raw],
                facecolors="none",
                edgecolors=COLORS[site],
                marker="s",
                s=60,
                linewidth=1.5,
            )
        ax.set(title=title, xlabel="Federated-round progress (%)")
        ax.set_xticks((20, 40, 60, 80, 100))
        if percentage:
            ax.set_ylabel("Percent")
    handles, labels = axes[0, 0].get_legend_handles_labels()
    fig.legend(
        handles,
        labels,
        loc="upper center",
        bbox_to_anchor=(0.5, 0.945),
        ncol=3,
        frameon=False,
    )
    fig.suptitle(
        "Federated checkpoint trajectories (open square: native raw round 10)",
        y=0.995,
    )
    fig.tight_layout(rect=(0, 0, 1, 0.89))
    fig.savefig(assets / "benchmark_trajectories.png", bbox_inches="tight")
    plt.close(fig)


def add_inline(paragraph: Any, text: str) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def markdown_table(lines: Iterable[str]) -> list[list[str]]:
    return [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]


def make_docx(report: Path, output: Path) -> None:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (
        ("Title", 22, "17365D"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 11, "2F5597"),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
    if "Figure Caption" not in [style.name for style in document.styles]:
        caption_style = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Arial"
        caption_style.font.size = Pt(9)
        caption_style.font.italic = True

    report_title = next(
        (
            line.lstrip("#").strip()
            for line in report.read_text(encoding="utf-8").splitlines()
            if line.startswith("# ")
        ),
        report.stem.replace("_", " "),
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{report_title}  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    lines = report.read_text(encoding="utf-8").splitlines()
    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        image = re.fullmatch(r"!\[(.+)]\((.+)\)", line)
        if image:
            image_path = (report.parent / image.group(2)).resolve()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(image_path), width=Inches(6.45))
            caption = document.add_paragraph(style="Figure Caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run(image.group(1))
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = markdown_table(table_lines)
            if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
                rows.pop(1)
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            table.autofit = True
            for row_index, values in enumerate(rows):
                for column_index, value in enumerate(values):
                    cell = table.cell(row_index, column_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.text = ""
                    add_inline(cell.paragraphs[0], value)
                    for run in cell.paragraphs[0].runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(7.5 if len(values) >= 7 else 8.5)
                        if row_index == 0:
                            run.bold = True
                    if row_index == 0:
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "D9EAF7")
                        cell._tc.get_or_add_tcPr().append(shading)
            document.add_paragraph().paragraph_format.space_after = Pt(1)
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            if level == 1 and first_heading:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(text)
                first_heading = False
            else:
                document.add_heading(text, level=min(level, 3))
            index += 1
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, line[2:])
            index += 1
            continue
        paragraph = document.add_paragraph()
        add_inline(paragraph, line)
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--run-root", type=Path)
    parser.add_argument("--report", type=Path, required=True)
    parser.add_argument("--assets", type=Path, required=True)
    parser.add_argument("--docx", type=Path)
    parser.add_argument("--plots", action="store_true")
    args = parser.parse_args()
    if args.plots:
        if args.run_root is None:
            parser.error("--run-root is required with --plots")
        make_plots(args.run_root.resolve(), args.assets.resolve())
    if args.docx:
        make_docx(args.report.resolve(), args.docx.resolve())
    if not args.plots and not args.docx:
        parser.error("select --plots and/or --docx")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
